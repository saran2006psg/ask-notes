import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
import docx

logger = logging.getLogger(__name__)


def _save_docx_images(file_path: Path, images_dir: Path) -> List[Dict[str, str]]:
    """
    Extract all inline images from a DOCX file's relationships.
    Saves them to images_dir and returns a list of {path, filename} dicts.
    """
    import zipfile
    import io
    saved = []
    images_dir.mkdir(parents=True, exist_ok=True)

    try:
        with zipfile.ZipFile(file_path, "r") as z:
            img_names = [n for n in z.namelist() if n.startswith("word/media/")]
            for idx, img_name in enumerate(img_names):
                try:
                    img_data = z.read(img_name)
                    if len(img_data) < 500:
                        continue  # Skip tiny images

                    ext = Path(img_name).suffix.lower().lstrip(".")
                    if ext not in ("png", "jpg", "jpeg", "gif", "bmp"):
                        ext = "png"

                    # Normalise to PNG via Pillow
                    try:
                        from PIL import Image
                        img = Image.open(io.BytesIO(img_data))
                        if img.mode not in ("RGB", "RGBA", "L"):
                            img = img.convert("RGB")
                        out_buf = io.BytesIO()
                        img.save(out_buf, format="PNG")
                        img_data = out_buf.getvalue()
                        ext = "png"
                    except Exception:
                        pass

                    img_filename = f"img_{idx + 1}.{ext}"
                    img_path = images_dir / img_filename
                    with open(img_path, "wb") as f:
                        f.write(img_data)

                    saved.append({"path": str(img_path), "filename": img_filename, "page": 1})
                    logger.debug(f"Saved DOCX image: {img_filename}")
                except Exception as e:
                    logger.warning(f"Failed to save DOCX image {img_name}: {e}")
    except Exception as e:
        logger.warning(f"Failed to extract images from DOCX {file_path.name}: {e}")

    return saved


def extract_docx_data(file_path: Path, images_dir: Optional[Path] = None) -> dict:
    """
    Extracts text from a DOCX file. Since DOCX doesn't have hard physical pages easily
    accessible, we group paragraphs and tables into virtual 'pages' (approx 400 words or split by page breaks).
    Optionally extracts inline images if images_dir is provided.

    Returns:
        dict: {
            "filename": str,
            "total_pages": int,
            "pages": [
                {
                    "page_number": int,
                    "text": str,
                    "images": [...]   # if images extracted
                }
            ]
        }
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    logger.info(f"Starting text extraction for: {file_path.name}")

    # Extract all DOCX images upfront if images_dir provided
    all_doc_images: List[Dict[str, str]] = []
    if images_dir:
        all_doc_images = _save_docx_images(file_path, images_dir)

    try:
        doc = docx.Document(file_path)
        pages_data = []
        current_page_text = []
        current_word_count = 0
        page_num = 1

        def flush_page():
            nonlocal page_num, current_page_text, current_word_count
            text = "\n".join(current_page_text).strip()
            if text:
                # Attach all doc images to first page (DOCX doesn't map images to pages)
                page_images = all_doc_images if page_num == 1 else []
                pages_data.append({
                    "page_number": page_num,
                    "text": text,
                    "images": page_images
                })
                page_num += 1
            current_page_text = []
            current_word_count = 0

        # Iterate paragraphs and check for page breaks
        for para in doc.paragraphs:
            has_page_break = False
            for run in para.runs:
                # Check for page break element or last rendered page break
                if 'w:br' in run._r.xml and 'type="page"' in run._r.xml:
                    has_page_break = True
                elif 'w:lastRenderedPageBreak' in run._r.xml:
                    has_page_break = True
                    
            para_text = para.text.strip()
            if para_text:
                current_page_text.append(para_text)
                current_word_count += len(para_text.split())
                
            if has_page_break or current_word_count >= 400:
                flush_page()
                
        # Process tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                # Deduplicate side-by-side cell texts
                row_cells = []
                for cell in row.cells:
                    cell_val = cell.text.strip()
                    if cell_val and (not row_cells or row_cells[-1] != cell_val):
                        row_cells.append(cell_val)
                if row_cells:
                    table_text.append(" | ".join(row_cells))
            
            if table_text:
                current_page_text.append("[Table]\n" + "\n".join(table_text))
                current_word_count += len("\n".join(table_text).split())
                
            if current_word_count >= 400:
                flush_page()
                
        # Flush remaining content
        flush_page()
        
        # Fallback if doc is empty
        if not pages_data:
            pages_data.append({
                "page_number": 1,
                "text": "",
                "images": all_doc_images
            })

        logger.info(f"Successfully extracted {len(pages_data)} virtual pages from {file_path.name}")
        return {
            "filename": file_path.name,
            "total_pages": len(pages_data),
            "pages": pages_data
        }
    except Exception as e:
        logger.error(f"Error reading DOCX {file_path.name}: {e}")
        raise e
